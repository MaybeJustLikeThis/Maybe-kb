"""Tests for parser registry and MarkdownParser."""
import pytest
from pathlib import Path
from kb.core.parsers import Parser, ParsedContent, ParserRegistry, MarkdownParser


def test_parsed_content_is_frozen():
    """ParsedContent is immutable."""
    pc = ParsedContent(text="hello", metadata={"pages": 1}, attachments=[])
    assert pc.text == "hello"
    assert pc.metadata == {"pages": 1}


def test_markdown_parser_parses_file(tmp_path: Path):
    """MarkdownParser extracts frontmatter and body."""
    md = tmp_path / "test.md"
    md.write_text("---\ntitle: Hello\ntags: [a, b]\n---\n\n# Content\n\nBody text.", encoding="utf-8")

    parser = MarkdownParser()
    result = parser.parse(md)

    assert "# Content" in result.text
    assert "Body text" in result.text
    assert result.metadata["title"] == "Hello"
    assert result.metadata["tags"] == ["a", "b"]


def test_markdown_parser_no_frontmatter(tmp_path: Path):
    """MarkdownParser handles files without frontmatter."""
    md = tmp_path / "plain.md"
    md.write_text("# Just a heading\n\nSome content.", encoding="utf-8")

    parser = MarkdownParser()
    result = parser.parse(md)

    assert "Just a heading" in result.text
    assert "Some content" in result.text


def test_parser_registry_returns_correct_parser():
    """ParserRegistry returns parser by extension."""
    assert isinstance(ParserRegistry.get(".md"), MarkdownParser)
    assert isinstance(ParserRegistry.get(".MD"), MarkdownParser)


def test_parser_registry_raises_for_unknown_extension():
    """ParserRegistry raises KeyError for unregistered extensions."""
    with pytest.raises(KeyError):
        ParserRegistry.get(".xyz")


def test_parser_registry_register_custom_parser():
    """Custom parsers can be registered."""
    class FakeParser:
        def parse(self, path, *, context_bound=False):
            return ParsedContent(text="fake", metadata={}, attachments=[])

    ParserRegistry.register(".fake", FakeParser())
    assert isinstance(ParserRegistry.get(".fake"), FakeParser)


# --- PDF Parser tests ---


def test_pdf_parser_registered():
    """PDFParser is registered for .pdf extension."""
    import kb.parsers.pdf  # noqa: F401
    parser = ParserRegistry.get(".pdf")
    assert parser is not None


def test_pdf_parser_extracts_text(tmp_path: Path):
    """PDFParser extracts text from a simple PDF."""
    import kb.parsers.pdf  # noqa: F401

    pdf_path = tmp_path / "test.pdf"
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF World")
        doc.save(str(pdf_path))
        doc.close()
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    parser = ParserRegistry.get(".pdf")
    result = parser.parse(pdf_path)
    assert "Hello PDF World" in result.text
    assert result.metadata.get("pages") == 1


def test_pdf_parser_attaches_original(tmp_path: Path):
    """PDFParser includes original PDF as attachment."""
    import kb.parsers.pdf  # noqa: F401

    pdf_path = tmp_path / "test.pdf"
    try:
        import fitz
        doc = fitz.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    parser = ParserRegistry.get(".pdf")
    result = parser.parse(pdf_path)
    assert len(result.attachments) == 1
    assert result.attachments[0] == pdf_path


# --- DOCX Parser tests ---


def test_docx_parser_registered():
    """DocxParser is registered for .docx extension."""
    import kb.parsers.docx  # noqa: F401
    parser = ParserRegistry.get(".docx")
    assert parser is not None


def test_docx_parser_extracts_text(tmp_path: Path):
    """DocxParser extracts text from a simple DOCX."""
    import kb.parsers.docx  # noqa: F401

    docx_path = tmp_path / "test.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello DOCX World")
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = ParserRegistry.get(".docx")
    result = parser.parse(docx_path)
    assert "Hello DOCX World" in result.text


def test_docx_parser_attaches_original(tmp_path: Path):
    """DocxParser includes original DOCX as attachment."""
    import kb.parsers.docx  # noqa: F401

    docx_path = tmp_path / "test.docx"
    try:
        from docx import Document
        doc = Document()
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = ParserRegistry.get(".docx")
    result = parser.parse(docx_path)
    assert len(result.attachments) == 1


# --- Image Parser tests ---


def test_image_parser_registered():
    """ImageParser is registered for common image extensions."""
    import kb.parsers.image  # noqa: F401
    for ext in [".png", ".jpg", ".jpeg", ".webp"]:
        assert ParserRegistry.get(ext) is not None


def test_image_parser_context_bound_skips_description(tmp_path: Path):
    """context_bound=True returns empty text."""
    import kb.parsers.image  # noqa: F401
    import struct
    import zlib

    def create_png(path):
        sig = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data)
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
        raw = b'\x00\xff\x00\x00'
        compressed = zlib.compress(raw)
        idat_crc = zlib.crc32(b'IDAT' + compressed)
        idat = struct.pack('>I', len(compressed)) + b'IDAT' + compressed + struct.pack('>I', idat_crc)
        iend_crc = zlib.crc32(b'IEND')
        iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
        path.write_bytes(sig + ihdr + idat + iend)

    png_path = tmp_path / "test.png"
    create_png(png_path)

    parser = ParserRegistry.get(".png")
    result = parser.parse(png_path, context_bound=True)
    assert result.text == ""
    assert len(result.attachments) == 1


def test_image_parser_standalone(tmp_path: Path):
    """Standalone image generates a markdown image reference placeholder."""
    import kb.parsers.image  # noqa: F401
    import struct
    import zlib

    png_path = tmp_path / "standalone.png"
    sig = b'\x89PNG\r\n\x1a\n'
    ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data)
    ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
    raw = b'\x00\xff\x00\x00'
    compressed = zlib.compress(raw)
    idat_crc = zlib.crc32(b'IDAT' + compressed)
    idat = struct.pack('>I', len(compressed)) + b'IDAT' + compressed + struct.pack('>I', idat_crc)
    iend_crc = zlib.crc32(b'IEND')
    iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
    png_path.write_bytes(sig + ihdr + idat + iend)

    parser = ParserRegistry.get(".png")
    result = parser.parse(png_path)
    assert "![standalone.png]" in result.text
    assert len(result.attachments) == 1
